from docx               import Document
from docx.shared        import Pt, RGBColor, Inches, Cm
from docx.enum.text     import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE, WD_LINE_SPACING
from docx.enum.style    import WD_STYLE_TYPE
from docx.enum.table     import WD_TABLE_ALIGNMENT


# *********************************************************
#                      PAGE 1 FUNCTIONS
# *********************************************************
def addTitle_P1(document):
    title = document.add_heading('Koala')
    title.style.font.name = 'Arial'

    #print(dir(title.style.font))

    title.style.font.size = Pt(26)
    #print(title.style.font.bold)
    title.style.font.bold = False
    #title.style.font.italic = True
    #title.style.font.underline = True

    #print(dir(title.style.font.color))

    title.style.font.color.rgb = RGBColor(46, 116, 181)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #print(dir(title.paragraph_format))
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(0)

    ...
# END def addTitle_P1

def addParagraph_P1(document):
    
    # EMPTY LINE PARAGRAPH
    emptyLinePara = document.add_paragraph()
    emptyLinePara.style.font.name = 'calibri (Body)'
    emptyLinePara.style.font.size = Pt(11)

    emptyLineParaFormat = emptyLinePara.paragraph_format
    emptyLineParaFormat.space_before = Pt(0)
    emptyLineParaFormat.space_after = Pt(8)

    # TEXT PARAGRAPHS
    paragraphTextList = []
    # with open(r'resources\text.txt') as fileStream:
    with open('resources/text.txt', encoding='utf8') as fileStream:
        paragraphTextList = fileStream.readlines()

    for idx, paragraphText in enumerate(paragraphTextList):
        paragraphTextList[idx] = paragraphText.replace('\n', '')

    # print(dir(document.styles))
    # print(help(document.styles.add_style))

    myStyle = document.styles.add_style('MyStyle', WD_STYLE_TYPE.PARAGRAPH)
    myStyle.font.name = 'Arial'
    myStyle.font.size = Pt(10.5)

    formatMyStyle = myStyle.paragraph_format
    formatMyStyle.space_before = Pt(0)
    formatMyStyle.space_after = Pt(8)
    formatMyStyle.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # PARAGRAPH 1
    para = document.add_paragraph('', myStyle)
    
    # print(dir(para.paragraph_format))
    paraFormart = para.paragraph_format
    paraFormart.line_spacing = 1.5

    paraBlock = paragraphTextList[0].split('#')

    for idx, paraBlock in enumerate(paraBlock):
        run = para.add_run(paraBlock)

        indices = [1, 3, 5, 7, 9, 11]

        if idx in indices:

            if idx in [1, 3, 5]:
                run.bold = True

            if idx in [5]:
                run.italic = True

            if idx in [9]:
                run.underline = True

            if idx in [11]:
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.underline = WD_UNDERLINE.DASH                
    # END for idx, paraBlock


    # PARAGRAPH 2
    para = document.add_paragraph(paragraphTextList[1], myStyle)
    
    # EMPTY LINE PARAGRAPH
    emptyLinePara = document.add_paragraph('', style=myStyle)

    # emptyLineParaFormat = emptyLinePara.paragraph_format
    # emptyLineParaFormat.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    # emptyLineParaFormat.line_spacing = 1.08
    
    ...
# END def addParagraph_P1

def addKoalaImage_P1(document):
    # print(dir(document))
    # print(help(document.add_picture))

    imgPath = 'resources\koalaOriginal-Copia.jpg'
    imgW = Inches(5.79)
    imgH = Inches(5.69)

    document.add_picture(imgPath, width=imgW, height=imgH)
    document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ...
# END def addKoalaImage_P1


# *********************************************************
#                      PAGE 2 FUNCTIONS
# *********************************************************

data = {
    "Title": "Adicionando uma tabela ao documento e usando uma base de dados",
    "Content": [
        {
            "Chapter": "Chapter 1 Text",
            "Text": "This is my text 1 \n This is still my text 1",
            "Image": "resources\koalaOriginal-Copia.jpg",
            "Table": ["A", "B", "C"]
        },
        {
            "Chapter": "Chapter 2 Text",
            "Text": "This is my text 2 \n This is still my text 2",
            "Image": "resources\koalaOriginal-Copia.jpg",
            "Table": ["1", "2", "3"]
        },
        {
            "Chapter": "Chapter 3 Text",
            "Text": "This is my text 3 \n This is still my text 3",
            "Image": "resources\koalaOriginal-Copia.jpg",
            "Table": ["1", "2", "3"]
        },
    ]
}

def addTitle_P2(document):
    title2 = document.add_heading(data.get("Title"), level=1)
    title2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for content in data.get("Content"):
        # Add paragraph with h2
        document.add_paragraph(content.get("Chapter"))
        # Add paragraph 
        document.add_paragraph(content.get("Text"))
        # Add Image
        document.add_picture(content.get("Image"), width=Inches(1.25))
        # Add Table 
        document.add_heading("Table", level=3)
        table = document.add_table(rows=1, cols=3, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Col_1"
        hdr_cells[1].text = "Col_2"
        hdr_cells[2].text = "Col_3"
        row_cells = table.add_row().cells

        for idx, element in enumerate(content.get("Table")):
            row_cells[idx].text = element
        



# ...
# END def addTitle_P2

# def addTable_P2(document):
#     ...
# END def addTable_P2




# *********************************************************
#                      DOCUMENTO
# *********************************************************
document = Document()


'''
Precisamos padronizar as margens do documento
'''

# Precisamos investigar as propriedades do document
#print(dir(document))


# Vamos veriricar a propriedade settings do document
# print(dir(document.settings))  não encontramos nada

# Vamos investigar as propriedades de sections do documento
# print(dir(document.sections))

section = document.sections[0]

# section.page_width = Pt(8.27)
# section.page_height = Pt(11.69)

section.page_width = Inches(8.27)
section.page_height = Inches(11.69)

section.bottom_margin = Inches(1)
section.top_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

#print(dir(section))
section.header_distance = Inches(0.49)
section.footer_distance = Inches(0.49)

# PAGE 1
addTitle_P1(document)
addParagraph_P1(document)
addKoalaImage_P1(document)

# PAGE BREAK

# PAGE 2 
addTitle_P2(document)
# addTable_P2(document)

document.save('koala.docx')